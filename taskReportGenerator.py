# -*- coding: utf-8 -*-
# toDate is exclusive!

import datetime
import shlex
import subprocess
import re
import os
import json

import docx

SVN_LOG_ENTRIES_SEPARATOR = '------------------------------------------------------------------------'
REPORT_FOLDER_NAME = 'reports'
SUBCONTENT_ENTRY_COUNT = 2

class DateWalker(object):
    def __init__(self, fromDate, toDate, delta):
        self.fromDate = fromDate
        self.toDate = toDate
        self.delta = delta
        
    def walk(self, dateVisitor):
        lowerDate = datetime.date.fromordinal(self.fromDate.toordinal())
        while True:
            upperDate = lowerDate + self.delta
            if upperDate > self.toDate:
                break
            dateVisitor.visitInterval((lowerDate, upperDate))
            lowerDate = upperDate
        if lowerDate < self.toDate:
            dateVisitor.visitInterval((lowerDate, self.toDate))
            
def issueCommand(command):
    print 'issue', command
    if isinstance(command, unicode):
        command = command.encode('utf-8')
    arguments = shlex.split(command)
    return subprocess.check_output(arguments)
    
def sanitizeEntry(entry):
    entry = re.sub('\n{2,}', '\n', entry) # remove unnecessary newlines
    for dummy in xrange(2): # remove first two lines
        entry = entry[entry.find('\n') + 1:]
    return entry
            
def querySVNLog(URL, user, password, fromDate, toDate):
    toDate = toDate - datetime.timedelta(days = 1)
    log = issueCommand('svn log "%s" --username "%s" --password %s -r {%s}:{"%s 23:59"}' % (URL, user, password, fromDate.isoformat(), toDate.isoformat()))
    entries = log.split(SVN_LOG_ENTRIES_SEPARATOR)
    userPattern = '\| %s \|' % (user,)
    if isinstance(userPattern, unicode):
        userPattern = userPattern.encode('utf-8') # compare with identical encoding
    entries = [entry for entry in entries if entry and re.search(userPattern, entry, re.IGNORECASE) is not None]
    if len(entries) > 0:
        # remove first entry if it falls outside the range
        match = re.search('(\\d{4})-(\\d{2})-(\\d{2})', entries[0])
        if match is not None:
            entryDate = datetime.date(int(match.group(1)), int(match.group(2)), int(match.group(3)))
            if entryDate < fromDate:
                entries = entries[1:]
    entries = [sanitizeEntry(entry) for entry in entries]
    return entries
    
def combineWordDocuments(filePaths, combinedDocumentName):
    if len(filePaths) < 2:
        return
        
    combinedDocument = docx.Document(filePaths[0])
    combinedDocument.add_page_break()
    filePaths = filePaths[1:]
    for index, path in enumerate(filePaths):
        doc = docx.Document(path)
        if index < len(filePaths) - 1:
            doc.add_page_break()
        for element in doc.part.element:
            combinedDocument.part.element.append(element)
    combinedDocument.save('%s.docx' % (combinedDocumentName,))
    
class DateVisitor(object):
    def visitInterval(self, interval):
        pass
    
class TaskReportGenerator(DateVisitor):
    def __init__(self):
        self.gotTimeIntervals = []
        
    def visitInterval(self, interval):
        print 'from %s to %s' % (interval[0].isoformat(), interval[1].isoformat())
        self.gotTimeIntervals.append(interval)
        
    def generateReports(self, SVNAccount, reportInfo):
        reportPaths = []
        for index, timeInterval in enumerate(self.gotTimeIntervals):
            mainContent = TaskReportGenerator.reportContentWithSVNLogEntryCount(SVNAccount['URLS'], SVNAccount['USER'], SVNAccount['PASSWORD'], timeInterval[0], timeInterval[1])
            subContent = u''
            if index < len(self.gotTimeIntervals) - 1:
                nextTimeInterval = self.gotTimeIntervals[index + 1]
                subContent = TaskReportGenerator.reportContentWithSVNLogEntryCount(SVNAccount['URLS'], SVNAccount['USER'], SVNAccount['PASSWORD'], nextTimeInterval[0], nextTimeInterval[1], SUBCONTENT_ENTRY_COUNT)
            path = self.produceReportDocument(reportInfo['TEMPLATE_PATH'], reportInfo['AUTHOR'], timeInterval[0], timeInterval[1] - datetime.timedelta(days = 1), mainContent, subContent)
            reportPaths.append(path)
        return reportPaths
        
    def produceReportDocument(self, template, name, fromDate, toDate, mainContent, subContent):
        """this method is template-specific"""
        reportDocument = docx.Document(template)
        propertiesRow = reportDocument.tables[0].rows[0]
        propertiesRow.cells[3].text = name
        propertiesRow.cells[5].text = fromDate.isoformat()
        propertiesRow.cells[7].text = toDate.isoformat()
        innerTable = reportDocument.tables[0].cell(1, 0).tables[0]
        if mainContent:
            innerTable.cell(0, 1).text = mainContent
        if subContent:
            innerTable.cell(1, 1).text = subContent
        reportDocumentPath = '%s.docx' % ('_'.join([name, fromDate.isoformat(), toDate.isoformat()]),)
        reportDocumentPath = os.path.join(REPORT_FOLDER_NAME, reportDocumentPath)
        if not os.path.exists(REPORT_FOLDER_NAME):
            os.makedirs(REPORT_FOLDER_NAME)
        reportDocument.save(reportDocumentPath)
        return reportDocumentPath
    
    @staticmethod
    def reportContentWithSVNLogEntryCount(SVNURLs, user, password, fromDate, toDate, entryCount = -1):
        if 0 == entryCount:
            return u''
        
        content = u''
        currentEntryCount = 0
        entryDelimiter = '-' * 20 + '\n'
        sectionDelimiter = '\n' + '=' * 30 + '\n'
        for URL in SVNURLs:
            logEntries = querySVNLog(URL, user, password, fromDate, toDate)
            if len(logEntries) < 1:
                continue
            if content:
                content += '\n'
            content += URL + sectionDelimiter
            if entryCount > 0:
                logEntries = logEntries[:min(entryCount - currentEntryCount, len(logEntries))]
            content += entryDelimiter.join(logEntries).decode('utf-8')
            currentEntryCount += len(logEntries)
            if entryCount > 0 and currentEntryCount >= entryCount:
                break
        return content
        
def dateFromDateString(dateString):
    components = [int(c) for c in dateString.split('-')]
    return datetime.date(*components)
    
def main():
    pathPrefix = u'./'
    configFileNameList = [fileEntry for fileEntry in os.listdir(pathPrefix) if os.path.isfile(os.path.join(pathPrefix, fileEntry))]
    configFileNameList = [configFileName for configFileName in configFileNameList if u'config_' in configFileName]
    for configFileName in configFileNameList:
        print 'processing', configFileName
        configFile = open(configFileName, 'r')
        config = None
        try:
            config = json.load(configFile)
        except:
            pass
        configFile.close()
        requiredKeySet = {'SVN_ACCOUNT', 'REPORT_INFO', 'TIME_DELTA', 'FROM_DATE', 'TO_DATE', 'OUTPUT_PATH'}
        if config is None or not set(config.keys()).issuperset(requiredKeySet):
            print 'could not generate report with %s' % (configFileName,)
            continue
        SVNAccount = config['SVN_ACCOUNT']
        reportInfo = config['REPORT_INFO']
        
        timeDelta = datetime.timedelta(days = int(config['TIME_DELTA']))
        fromDate = dateFromDateString(config['FROM_DATE'])
        toDate = dateFromDateString(config['TO_DATE'])
        walker = DateWalker(fromDate, toDate, timeDelta)
        generator = TaskReportGenerator()
        walker.walk(generator)
        reportPaths = generator.generateReports(SVNAccount, reportInfo)
        combineWordDocuments(reportPaths, config['OUTPUT_PATH'])

if '__main__' == __name__:
    main()
    